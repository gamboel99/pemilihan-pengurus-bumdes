import streamlit as st
import pandas as pd
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import qrcode
from io import BytesIO
from PIL import Image

# ======================= Fungsi Export Word =======================
def generate_word_with_ttd(data, penilai_identitas):
    doc = Document()

    # Kop Surat
    table = doc.add_table(rows=1, cols=3)
    row = table.rows[0].cells

    try:
        row[0].paragraphs[0].add_run().add_picture("assets/logo_pemdes.png", width=Inches(1))
    except:
        row[0].text = "Logo Pemdes"
    row[1].text = "PEMERINTAH DESA KELING\nBUMDes BUWANA RAHARJA\nKecamatan Kepung, Kabupaten Kediri"
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        row[2].paragraphs[0].add_run().add_picture("assets/logo_bumdes.png", width=Inches(1))
    except:
        row[2].text = "Logo BUMDes"

    doc.add_paragraph("")
    title = doc.add_paragraph("REKAPITULASI HASIL PENILAIAN CALON PENGURUS BUMDES")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    doc.add_paragraph("")

    # Tabel Nilai
    table = doc.add_table(rows=1, cols=len(data.columns))
    table.style = 'Table Grid'
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(data.columns):
        hdr_cells[i].text = col

    for _, row_data in data.iterrows():
        row = table.add_row().cells
        for i, col in enumerate(data.columns):
            row[i].text = str(round(row_data[col], 2)) if isinstance(row_data[col], float) else str(row_data[col])

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(12)

    doc.add_paragraph("")
    ttd = doc.add_table(rows=2, cols=2)
    ttd.style = "Table Grid"
    ttd.cell(0,0).text = "Mengetahui,\nPanitia Pemilihan"
    ttd.cell(0,1).text = f"Kediri, {datetime.now().strftime('%d %B %Y')}\nPenilai"
    ttd.cell(1,0).text = "(...........................)"
    ttd.cell(1,1).text = f"({penilai_identitas.get('nama','...')})"

    for row in ttd.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)

    doc.add_paragraph("")
    doc.add_paragraph("Legalitas:")
    doc.add_paragraph("Dokumen ini resmi diterbitkan oleh Panitia Pemilihan Pengurus\n"
                      "BUMDes Buwana Raharja Desa Keling, Kec. Kepung, Kab. Kediri.")

    qr = qrcode.make("Dokumen Resmi Panitia Pemilihan BUMDes Desa Keling")
    qr_io = BytesIO()
    qr.save(qr_io)
    qr_io.seek(0)
    doc.add_picture(qr_io, width=Inches(1.2))

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ======================= Konfigurasi App =======================
st.set_page_config(page_title="Polling Pengurus BUMDes", layout="wide")
st.title("📊 Sistem Polling Pengurus BUMDes Buwana Raharja")

DATA_FOLDER = "data"
PENILAI_FILE = os.path.join(DATA_FOLDER, "penilai.csv")
KANDIDAT_FILE = os.path.join(DATA_FOLDER, "kandidat.csv")
HASIL_FILE = os.path.join(DATA_FOLDER, "hasil_penilaian.csv")
os.makedirs(DATA_FOLDER, exist_ok=True)

bobot = {
    "Tes Psikologi": 0.15,
    "Tes MS Office": 0.15,
    "Presentasi Gagasan": 0.30,
    "Esai Refleksi Diri": 0.20,
    "Wawancara Panel": 0.20
}

if not os.path.exists(PENILAI_FILE):
    pd.DataFrame({"Nama Penilai": [
        "Kepala Desa", "Sekretaris Desa", "Ketua BPD", "Anggota BPD",
        "Kasi PMD Kecamatan", "Pendamping Kecamatan", "DPMPD"
    ]}).to_csv(PENILAI_FILE, index=False)

if not os.path.exists(KANDIDAT_FILE):
    pd.DataFrame({
        "Nama": [
            "Eko Wahyu Diantoro", "Racal Pudjo Alberto",
            "Nanda Rafiatul", "Wahida Nayla", "Fenny Alvionita",
            "Nani NurMahmudah", "Novia Lestari",
            "Jelvi Tri K", "Elton Priloro", "Azuma Zundana"
        ],
        "Posisi": [
            "Direktur Utama", "Direktur Utama",
            "Sekretaris", "Sekretaris", "Sekretaris",
            "Sekretaris", "Sekretaris",
            "Bendahara", "Bendahara", "Bendahara"
        ]
    }).to_csv(KANDIDAT_FILE, index=False)

if not os.path.exists(HASIL_FILE):
    pd.DataFrame(columns=["Penilai", "Posisi", "Nama"] + list(bobot.keys()) + ["Timestamp"]).to_csv(HASIL_FILE, index=False)

penilai_df = pd.read_csv(PENILAI_FILE)
kandidat_df = pd.read_csv(KANDIDAT_FILE)
hasil_df = pd.read_csv(HASIL_FILE)

# ======================= Form Identitas =======================
st.subheader("🧾 Identitas Penilai")
with st.form("form_identitas"):
    nama_penilai_form = st.text_input("Nama Lengkap Penilai")
    jabatan_penilai = st.text_input("Jabatan")
    instansi_penilai = st.text_input("Instansi")
    submitted_id = st.form_submit_button("✔️ Simpan Identitas")
    if submitted_id:
        st.session_state["identitas_penilai"] = {
            "nama": nama_penilai_form,
            "jabatan": jabatan_penilai,
            "instansi": instansi_penilai
        }
        st.success("✅ Identitas disimpan.")

# ======================= Form Penilaian =======================
st.subheader("📝 Form Penilaian")
if "identitas_penilai" in st.session_state:
    nama_penilai = st.session_state["identitas_penilai"]["nama"]
    posisi = st.selectbox("Pilih Posisi yang Dinilai:", kandidat_df["Posisi"].unique())
    kandidat_list = kandidat_df[kandidat_df["Posisi"] == posisi]["Nama"].tolist()

    with st.form("form_penilaian"):
        kandidat_dipilih = st.selectbox("Pilih Kandidat:", kandidat_list)

        sudah_nilai = (
            (hasil_df["Penilai"] == nama_penilai)
            & (hasil_df["Posisi"] == posisi)
            & (hasil_df["Nama"] == kandidat_dipilih)
        ).any()

        nilai_input = {}
        if sudah_nilai:
            st.warning("⚠️ Anda sudah menilai kandidat ini untuk posisi tersebut. Nilai tidak bisa diubah.")
        else:
            for aspek in bobot:
                nilai_input[aspek] = st.number_input(
                    f"{aspek} (0–100)", min_value=0, max_value=100, step=1,
                    key=f"{aspek}_{posisi}_{kandidat_dipilih}"
                )

        submitted = st.form_submit_button("💾 Simpan Penilaian")

        if submitted and not sudah_nilai:
            row = {
                "Penilai": nama_penilai,
                "Posisi": posisi,
                "Nama": kandidat_dipilih,
                "Timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            for aspek in bobot:
                row[aspek] = nilai_input[aspek]
            hasil_df = pd.concat([hasil_df, pd.DataFrame([row])], ignore_index=True)
            hasil_df.to_csv(HASIL_FILE, index=False)
            st.success("✅ Penilaian berhasil disimpan dan terkunci.")

# ======================= Rekap dan Export =======================
st.subheader("📈 Rekapitulasi Hasil Sementara")
if not hasil_df.empty:
    hasil_df["Total"] = hasil_df[[*bobot]].apply(
        lambda row: sum(row[aspek] * bobot[aspek] for aspek in bobot), axis=1)
    ranking_df = hasil_df.groupby(["Nama", "Posisi"]).agg({"Total": "mean"}).reset_index()
    ranking_df = ranking_df.sort_values(["Posisi", "Total"], ascending=[True, False])

    for p in ranking_df["Posisi"].unique():
        st.markdown(f"### 🏆 Hasil Sementara: {p}")
        st.dataframe(ranking_df[ranking_df["Posisi"] == p][["Nama", "Total"]].reset_index(drop=True))

    st.download_button(
        label="⬇️ Download Rekap Penilaian (.CSV)",
        data=ranking_df.to_csv(index=False).encode("utf-8"),
        file_name="rekap_penilaian.csv",
        mime="text/csv"
    )

    st.subheader("📤 Export Rekap ke Word (dengan TTD & Barcode)")
    if "identitas_penilai" in st.session_state:
        word_file = generate_word_with_ttd(ranking_df, st.session_state["identitas_penilai"])
        st.download_button(
            label="⬇️ Download Dokumen Word",
            data=word_file,
            file_name="Rekap_Penilaian_Pengurus_BUMDes.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ======================= Footer =======================
st.markdown("---")
st.markdown(
    "<div style='text-align: center;'>Developed by CV Mitra Utama Consultindo</div>",
    unsafe_allow_html=True
)
